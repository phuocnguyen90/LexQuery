# utils/file_handler.py

import json
import logging
from typing import List, Dict, Any, Optional, Union

import os
import pandas as pd
import tempfile
import zipfile
import subprocess

from docx import Document
from docxcompose.composer import Composer
from shared_libs.models.record_model import Record

# Configure logger
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def load_record(raw_input: str, llm_processor, is_formatted: bool = True) -> Optional[Record]:
    """
    Load a Record object from raw input, determining the format.

    :param raw_input: The raw input string (tagged text, JSON, or unformatted text).
    :param llm_processor: A callable that processes unformatted text.
    :param is_formatted: Boolean indicating if the input is formatted by default.
    :return: Record object or None if loading fails.
    """
    try:
        if is_formatted:
            # Attempt to parse as JSON
            try:
                record = Record.from_json(raw_input)
                if record:
                    logging.debug("Record loaded from JSON.")
                    return record
            except Exception as e:
                logging.debug(f"Failed to load record from JSON: {e}")

            # Attempt to parse as tagged text
            record = Record.from_tagged_text(raw_input)
            if record:
                logging.debug("Record loaded from tagged text.")
                return record

            # If parsing fails, treat as unformatted
            logging.debug("Record could not be parsed as formatted. Treating as unformatted.")
            record = Record.from_unformatted_text(raw_input, llm_processor)
            if record:
                logging.debug("Record loaded from unformatted text via LLM.")
                return record
        else:
            # Treat input as unformatted
            logging.debug("Input specified as unformatted. Processing with LLM.")
            record = Record.from_unformatted_text(raw_input, llm_processor)
            if record:
                logging.debug("Record loaded from unformatted text via LLM.")
                return record

        logging.error("Failed to load record from any supported format.")
        return None

    except Exception as e:
        logging.error(f"Unexpected error in load_record: {e}")
        return None

def read_input_file(file_path):
    """
    Read the raw input file and return its content.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = f.read()
        logging.info(f"Read data from '{file_path}'.")
        return data
    except Exception as e:
        logging.error(f"Error reading input file '{file_path}': {e}")
        raise



def write_output_file(file_path, data):
    """
    Write the processed data to the output file in JSON format.
    """
    try:
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logging.info(f"Wrote processed data to '{file_path}'.")
    except Exception as e:
        logging.error(f"Error writing to output file '{file_path}': {e}")


def save_processed_record(record: Dict[str, Any], file_path: str):
    """
    Save a single processed record to the output file.
    Overwrites if record with the same 'id' exists, or appends if it does not.
    
    :param record: Dictionary representing the processed record.
    :param file_path: Path to the output file.
    """
    output_2_jsonl(file_path, record)



def doc_to_docx_pipeline(input_doc_path, output_docx_path):

    """Pipeline to convert .doc to .docx and append OLE content."""

    # Step 1: Convert the original .doc file to .docx
    main_docx_content = doc_to_docx(input_doc_path)
    if main_docx_content is None:
        print(f"Failed to convert the original .doc file: {input_doc_path}")
        return

    # Save the converted main .docx to a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_main_docx:
        tmp_main_docx.write(main_docx_content)
        main_docx_path = tmp_main_docx.name

    # Step 2: Extract and process OLE objects
    ole_objects = extract_ole_objects(main_docx_path)

    # Convert and collect embedded .doc content to append
    appended_docs = []
    for i, ole_content in enumerate(ole_objects):
        # Convert the extracted .doc or .zip containing .doc files
        extracted_docs = process_ole_content(ole_content)
        for extracted_doc_content in extracted_docs:
            appended_docs.append(extracted_doc_content)

    # Step 3: Append the extracted contents to the main document
    if appended_docs:
        combined_doc = append_appendix_2_doc(main_docx_path, appended_docs)
        # Save the combined document to the specified output path
        combined_doc.save(output_docx_path)
        print(f"Combined document saved at: {output_docx_path}")
    else:
        print("No content to append.")
        # Save the original .docx content without changes
        with open(output_docx_path, 'wb') as f:
            f.write(main_docx_content)

def doc_to_docx(input_doc_path):

    """Convert a .doc file to .docx using LibreOffice."""
    with tempfile.TemporaryDirectory() as tmpdirname:
        # Copy the input .doc file to the temporary directory
        doc_path = os.path.join(tmpdirname, 'input.doc')
        with open(doc_path, 'wb') as doc_file:
            with open(input_doc_path, 'rb') as f:
                doc_file.write(f.read())

        # Convert the .doc file to .docx using LibreOffice
        try:
            subprocess.run(['soffice', '--headless', '--convert-to', 'docx', doc_path, '--outdir', tmpdirname],
                           check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE)
        except subprocess.CalledProcessError as e:
            print(f"Error during conversion: {e}")
            return None

        # Read the converted .docx content
        docx_path = os.path.join(tmpdirname, 'input.docx')
        if os.path.exists(docx_path):
            with open(docx_path, 'rb') as docx_file:
                return docx_file.read()
        else:
            print("Conversion failed: .docx file not found.")
            return None

def extract_ole_objects(docx_path):

    """Extract OLE objects from a .docx file."""
    ole_objects = []
    with zipfile.ZipFile(docx_path, 'r') as docx:
        # Locate embedded OLE objects
        for item in docx.namelist():
            if item.startswith('word/embeddings'):
                # Extract the OLE object file
                with docx.open(item) as file:
                    ole_objects.append(file.read())
    return ole_objects

def process_ole_content(ole_content):

    """Process OLE content, handle both .doc files and .zip containing .doc files."""

    processed_docs = []

    with tempfile.TemporaryDirectory() as tmpdirname:
        # Save the OLE content to a temporary file
        ole_path = os.path.join(tmpdirname, 'embedded_object')
        with open(ole_path, 'wb') as ole_file:
            ole_file.write(ole_content)

        # Try to handle OLE content as a .doc file directly
        try:
            # Convert the binary content directly using the helper function
            docx_content = doc_to_docx(ole_path)
            if docx_content:
                processed_docs.append(docx_content)
        except Exception as e:
            print(f"Failed to convert as a .doc file: {e}")

        # If the above fails, try to handle OLE content as a .zip file containing .doc files
        if not processed_docs and zipfile.is_zipfile(ole_path):
            try:
                with zipfile.ZipFile(ole_path, 'r') as zip_file:
                    for item in zip_file.namelist():
                        if item.endswith('.doc'):
                            with zip_file.open(item) as doc_file:
                                doc_content = doc_file.read()
                                # Save the content to a temporary .doc file
                                with tempfile.NamedTemporaryFile(delete=False, suffix='.doc') as temp_doc:
                                    temp_doc.write(doc_content)
                                    temp_doc.flush()
                                    # Convert the extracted .doc to .docx
                                    docx_content = doc_to_docx(temp_doc.name)
                                    if docx_content:
                                        processed_docs.append(docx_content)
            except Exception as e:
                print(f"Failed to process the OLE as a .zip file: {e}")

    return processed_docs


def append_appendix_2_doc(main_doc_path, appended_docs):
    """Append documents to the main document and return the combined document."""
    
    # Load the main document

    main_doc = Document(main_doc_path)
    composer = Composer(main_doc)

    # Append each document to the main document
    for doc_content in appended_docs:
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_docx:
            tmp_docx.write(doc_content)
            tmp_docx.flush()
            tmp_doc = Document(tmp_docx.name)
            composer.append(tmp_doc)

    return main_doc



def create_documents_dataframe(base_folder):
    """
    Traverse the folder structure and create a dataframe containing information
    about legal documents, including their hierarchy and related documents.

    Args:
    - base_folder: The root directory containing the documents, categorized by type.

    Returns:
    - A pandas DataFrame with columns: 'Category', 'Document ID', 'Hierarchy Level', 'Document Path', 'Parent Document ID'
    """
    documents_data = []

    # Traverse the base folder
    for category in os.listdir(base_folder):
        category_path = os.path.join(base_folder, category)

        if os.path.isdir(category_path):
            # Traverse files and subdirectories within the category
            for root, dirs, files in os.walk(category_path):
                for file in files:
                    # Extract the hierarchy identifier (e.g., "1.1.1") and the rest of the document name
                    file_parts = file.split(' ', 1)
                    if len(file_parts) < 2:
                        continue  # Skip files that don't match the expected naming convention

                    hierarchy_id = file_parts[0]  # "1.1.1" part
                    doc_name = file_parts[1]  # Remaining part of the document name
                    doc_id = hierarchy_id  # Document ID is the hierarchy identifier

                    # Determine hierarchy level based on the hierarchy_id
                    hierarchy_level = len(hierarchy_id.split('.'))  # e.g., "1.1.1" has 3 dots -> level 3

                    # Determine the parent document ID
                    parent_doc_id = None
                    if hierarchy_level > 1:  # Only look for a parent if it's not the top level
                        parent_hierarchy_id = '.'.join(hierarchy_id.split('.')[:-1])  # Remove the last part
                        parent_doc_id = parent_hierarchy_id  # Parent is the preceding identifier

                    # Add document info to the list
                    documents_data.append({
                        'Category': category,
                        'Document ID': doc_id,
                        'Document Name': doc_name,
                        'Hierarchy Level': hierarchy_level,
                        'Document Path': os.path.join(root, file),
                        'Parent Document ID': parent_doc_id
                    })

    # Create a DataFrame
    df = pd.DataFrame(documents_data)
    return df

def generate_unique_id(prefix: str = "REC") -> str:
    """
    Generate a unique ID for records without an existing ID.
    
    :param prefix: Prefix for the unique ID.
    :return: A unique ID string.
    """
    import uuid
    unique_part = uuid.uuid4().hex[:8].upper()
    return f"{prefix}_{unique_part}"




def output_2_jsonl(file_path: str, records: Union[Dict[str, Any], List[Union[Dict[str, Any], Any]]]):
    """
    Append processed record(s) to the output file in JSONL format.
    If a record with the same 'record_id' or 'id' exists, overwrite it.
    If the file does not exist, create it.

    Records are stored as JSON objects, one per line.

    :param file_path: Path to the output file.
    :param records: A single dictionary, a Record instance, or a list of them representing the processed record(s).
    """
    try:
        # Normalize records to a list of dicts
        records_to_add = []
        if isinstance(records, dict):
            records_to_add.append(records)
        elif isinstance(records, list):
            for record in records:
                if isinstance(record, dict):
                    records_to_add.append(record)
                elif hasattr(record, 'to_dict') and callable(getattr(record, 'to_dict')):
                    records_to_add.append(record.to_dict())
                else:
                    logger.warning("Skipping record that is neither a dict nor has a 'to_dict' method.")
        elif hasattr(records, 'to_dict') and callable(getattr(records, 'to_dict')):
            records_to_add.append(records.to_dict())
        else:
            logger.error("The 'records' parameter must be a dictionary, a Record instance with 'to_dict', or a list of them.")
            return

        # Initialize a dictionary to hold existing records for quick lookup
        existing_records_dict = {}

        # Check if the output file exists and load existing records
        if os.path.exists(file_path):
            logger.debug(f"Output file '{file_path}' exists. Reading existing records.")
            with open(file_path, 'r', encoding='utf-8') as f:
                for line in f:
                    line = line.strip()
                    if not line:
                        continue
                    try:
                        record = json.loads(line)
                        # Use 'record_id' as the primary identifier, fallback to 'id' if necessary
                        record_id = record.get('record_id') or record.get('id')
                        if record_id:
                            existing_records_dict[record_id] = record
                        else:
                            logger.warning("Existing record does not contain 'record_id' or 'id'. Skipping.")
                    except json.JSONDecodeError as e:
                        logger.error(f"JSON decoding error while reading '{file_path}': {e}")
                        continue
        else:
            logger.info(f"Output file '{file_path}' does not exist. It will be created.")

        # Process each new record
        for record in records_to_add:
            if not isinstance(record, dict):
                logger.warning("Skipping non-dictionary record.")
                continue

            # Use 'record_id' as the primary identifier, fallback to 'id' if necessary
            record_id = record.get('record_id') or record.get('id')
            if not record_id:
                logger.error("Record does not contain a 'record_id' or 'id' field. Skipping.")
                continue

            if record_id in existing_records_dict:
                logger.info(f"Overwriting existing record with ID: {record_id}.")
            else:
                logger.info(f"Appending new record with ID: {record_id}.")

            # Update or add the record in the existing_records_dict
            existing_records_dict[record_id] = record

        # Write all records back to the file in JSONL format
        with open(file_path, 'w', encoding='utf-8') as f:
            for rec_id, rec in existing_records_dict.items():
                json.dump(rec, f, ensure_ascii=False)
                f.write('\n')  # Newline separator between records

        logger.debug(f"Successfully saved {len(existing_records_dict)} record(s) to '{file_path}'.")

    except Exception as e:
        logger.error(f"An error occurred in output_to_jsonl: {e}")


def extract_text_from_txt(file_path):
    """
    Extracts text from a .txt file.
    
    Args:
        file_path (str): Path to the .txt file.
        
    Returns:
        str: Extracted text.
    """
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
        logger.info(f"Extracted text from TXT file: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading TXT file {file_path}: {e}")
        return ""

def extract_text_from_docx(file_path):
    """
    Extracts text from a .docx file, including text inside tables,
    preserving the order of paragraphs and tables as in the original document,
    and ensuring each table cell is on its own line.

    Args:
        file_path (str): Path to the .docx file.

    Returns:
        str: Extracted text.
    """
    import docx
    try:
        doc = docx.Document(file_path)
        full_text = []

        for block in iter_block_items(doc):
            if isinstance(block, docx.text.paragraph.Paragraph):
                if block.text.strip():
                    full_text.append(block.text.strip())
            elif isinstance(block, docx.table.Table):
                # Extract text from table
                full_text.append("<table>")
                for row in block.rows:
                    full_text.append("<tr>")
                    for cell in row.cells:
                        cell_text = cell.text.strip()
                        if cell_text:
                            # Add each cell content as a separate line
                            full_text.append(f"<td>{cell_text}</td>")
                    full_text.append("</tr>")
                full_text.append("</table>")

        text = '\n'.join(full_text)
        logger.info(f"Extracted text from DOCX file: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading DOCX file {file_path}: {e}")
        return ""


def iter_block_items(parent):
    """
    Generate a reference to each paragraph and table child within parent, in document order.

    Args:
        parent: The parent document or element.

    Yields:
        Each child paragraph and table in document order.
    """
    import docx
    from docx.oxml.ns import qn
    from docx.oxml.text.paragraph import CT_P
    from docx.oxml.table import CT_Tbl
    from docx.table import _Cell, Table
    from docx.text.paragraph import Paragraph

    if isinstance(parent, docx.document.Document):
        parent_element = parent.element.body
    elif isinstance(parent, _Cell):
        parent_element = parent._tc
    else:
        raise ValueError("Invalid parent type")

    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)



def extract_text_from_pdf(file_path):
    """
    Extracts text from a .pdf file.
    
    Args:
        file_path (str): Path to the .pdf file.
        
    Returns:
        str: Extracted text.
    """
    import PyPDF2

    try:
        reader = PyPDF2.PdfReader(file_path)
        text = ""
        for page_num in range(len(reader.pages)):
            page = reader.pages[page_num]
            extracted = page.extract_text()
            if extracted:
                text += extracted + '\n'
        logger.info(f"Extracted text from PDF file: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading PDF file {file_path}: {e}")
        return ""

def extract_text_from_html(file_path):
    """
    Extracts text from a .html file.
    
    Args:
        file_path (str): Path to the .html file.
        
    Returns:
        str: Extracted text.
    """
    from bs4 import BeautifulSoup
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            soup = BeautifulSoup(f, 'lxml')
        text = soup.get_text(separator='\n')
        logger.info(f"Extracted text from HTML file: {file_path}")
        return text
    except Exception as e:
        logger.error(f"Error reading HTML file {file_path}: {e}")
        return ""

def determine_file_type(file_path):
    """
    Determines the file type based on its extension.
    
    Args:
        file_path (str): Path to the file.
        
    Returns:
        str: File type ('txt', 'docx', 'pdf', 'html') or 'unsupported'.
    """
    _, ext = os.path.splitext(file_path)
    ext = ext.lower()
    if ext == '.txt':
        return 'txt'
    elif ext == '.docx':
        return 'docx'
    elif ext == '.pdf':
        return 'pdf'
    elif ext in ['.html', '.htm']:
        return 'html'
    else:
        return 'unsupported'